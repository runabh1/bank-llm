"""
Document Processor for Bank LLM System
Handles parsing PDF, DOCX, XLSX files and ingesting them into ChromaDB.
Extracts metadata: circular number, date, issuing authority.
"""

import re
import uuid
import logging
from pathlib import Path
from datetime import datetime
from typing import Optional

import chromadb
from chromadb.config import Settings

# Document parsers
import pypdf
import docx
import openpyxl

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings

# Configure pytesseract for OCR (Windows)
try:
    import pytesseract
    # Try to set the path for Windows installation
    pytesseract.pytesseract.pytesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
except ImportError:
    pytesseract = None

from config import (
    VECTOR_DB_DIR,
    CHROMA_COLLECTION_INTERNAL,
    EMBED_MODEL,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    OLLAMA_BASE_URL,
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ─── ChromaDB Client ──────────────────────────────────────────────────────────

def get_chroma_client():
    return chromadb.PersistentClient(
        path=str(VECTOR_DB_DIR),
        settings=Settings(anonymized_telemetry=False)
    )


def get_collection(collection_name: str):
    client = get_chroma_client()
    return client.get_or_create_collection(
        name=collection_name,
        metadata={"hnsw:space": "cosine"}
    )


# ─── Embedding Function ───────────────────────────────────────────────────────

def get_embeddings():
    return OllamaEmbeddings(
        model=EMBED_MODEL,
        base_url=OLLAMA_BASE_URL
    )


# ─── Metadata Extraction ──────────────────────────────────────────────────────

def extract_circular_metadata(text: str, filename: str) -> dict:
    """
    Try to extract circular reference number, date, and issuing authority
    from the first ~1000 characters of document text.
    """
    header_text = text[:1500]

    # Circular / Reference number patterns
    ref_patterns = [
        r'(?:Circular|Ref(?:erence)?|No\.?|Notification)[:\s#]*([A-Z0-9\.\-\/]+)',
        r'(?:DO|DOR|DBR|FIDD|DNBR|RBI)[\/\.\-][A-Z0-9]+[\/\.\-][A-Z0-9]+[\/\.\-]?\d*',
        r'(?:Circ(?:ular)?\.?\s*No\.?\s*)([A-Z0-9\-\/\.]+)',
        r'([A-Z]{2,}[\./][A-Z0-9]+[\./]\d{4}[\./]\d+)',
    ]
    ref_no = "N/A"
    for pat in ref_patterns:
        m = re.search(pat, header_text, re.IGNORECASE)
        if m:
            ref_no = m.group(0).strip()
            break

    # Date patterns
    date_patterns = [
        r'\b(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})\b',
        r'\b(\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})\b',
        r'\b((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})\b',
    ]
    date = "N/A"
    for pat in date_patterns:
        m = re.search(pat, header_text, re.IGNORECASE)
        if m:
            date = m.group(1).strip()
            break

    # Issuing authority
    authority = "Internal"
    authority_keywords = {
        "Reserve Bank of India": "RBI",
        "RBI": "RBI",
        "NABARD": "NABARD",
        "SEBI": "SEBI",
        "Assam Gramin Vikas Bank": "AGVB",
        "AGVB": "AGVB",
        "Pragati Gramin Bank": "PGB",
        "PGB": "PGB",
        "State Bank of India": "SBI",
        "Punjab National Bank": "PNB",
    }
    for keyword, short in authority_keywords.items():
        if keyword.lower() in header_text.lower():
            authority = short
            break

    return {
        "filename": filename,
        "ref_no": ref_no,
        "date": date,
        "authority": authority,
        "ingested_at": datetime.now().isoformat(),
    }


# ─── Document Parsers ─────────────────────────────────────────────────────────

def parse_pdf(filepath: Path) -> str:
    text = ""
    
    # Method 1: pypdf (fast, standard)
    try:
        with open(filepath, 'rb') as f:
            reader = pypdf.PdfReader(f)
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as page_err:
                    logger.warning(f"Error on page {i} of {filepath.name}: {page_err}")
    except Exception as e:
        logger.warning(f"pypdf parsing error {filepath.name}: {e}")

    # Method 2: PyMuPDF (fitz) - better for corrupted PDFs
    if len(text.strip()) < 50:
        try:
            import fitz
            logger.info(f"  Trying PyMuPDF (fitz) for {filepath.name}...")
            doc = fitz.open(str(filepath))
            fitz_text = ""
            for page_num in range(len(doc)):
                try:
                    page = doc[page_num]
                    page_text = page.get_text()
                    if page_text:
                        fitz_text += page_text + "\n"
                except Exception as page_err:
                    logger.warning(f"PyMuPDF page {page_num} error in {filepath.name}: {page_err}")
            doc.close()
            if fitz_text and len(fitz_text.strip()) > len(text.strip()):
                logger.info(f"  PyMuPDF extracted {len(fitz_text)} chars from {filepath.name}")
                text = fitz_text
        except ImportError:
            logger.warning("PyMuPDF (fitz) not installed — skipping")
        except Exception as e2:
            logger.warning(f"PyMuPDF parsing error {filepath.name}: {e2}")

    # Method 3: pdfminer fallback
    if len(text.strip()) < 50:
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            pdfminer_text = pdfminer_extract(str(filepath))
            if pdfminer_text and len(pdfminer_text.strip()) > len(text.strip()):
                logger.info(f"  pdfminer fallback yielded {len(pdfminer_text)} chars for {filepath.name}")
                text = pdfminer_text
        except ImportError:
            logger.warning("pdfminer.six not installed — skipping PDF fallback extraction")
        except Exception as e2:
            logger.warning(f"pdfminer fallback also failed for {filepath.name}: {e2}")

    # Method 4: OCR (pytesseract) for scanned/image-only PDFs
    if len(text.strip()) < 50:
        try:
            from pdf2image import convert_from_path
            import pytesseract
            logger.info(f"  Attempting OCR extraction for {filepath.name}...")
            images = convert_from_path(str(filepath))
            ocr_text = ""
            for i, image in enumerate(images):
                try:
                    page_ocr = pytesseract.image_to_string(image)
                    if page_ocr:
                        ocr_text += page_ocr + "\n"
                except Exception as ocr_err:
                    logger.warning(f"OCR failed on page {i} of {filepath.name}: {ocr_err}")
            if ocr_text and len(ocr_text.strip()) > len(text.strip()):
                logger.info(f"  OCR extraction yielded {len(ocr_text)} chars for {filepath.name}")
                text = ocr_text
        except ImportError as ie:
            logger.warning(f"OCR packages not installed (pdf2image or pytesseract) — cannot process scanned PDFs: {ie}")
        except Exception as e3:
            logger.warning(f"OCR extraction failed for {filepath.name}: {e3}")

    return text.strip()


def parse_txt(filepath: Path) -> str:
    """Parse a plain text file."""
    text = ""
    # Try common encodings
    for encoding in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'utf-16']:
        try:
            text = filepath.read_text(encoding=encoding)
            if text.strip():
                break
        except (UnicodeDecodeError, UnicodeError):
            continue
        except Exception as e:
            logger.error(f"TXT parse error {filepath.name} with {encoding}: {e}")
            break
    return text.strip()


def parse_docx(filepath: Path) -> str:
    text = ""
    try:
        doc = docx.Document(str(filepath))
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += row_text + "\n"
    except Exception as e:
        logger.error(f"DOCX parse error {filepath}: {e}")
    return text.strip()


def parse_xlsx(filepath: Path) -> str:
    text = ""
    try:
        wb = openpyxl.load_workbook(str(filepath), data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text += f"\n[Sheet: {sheet_name}]\n"
            for row in ws.iter_rows(values_only=True):
                row_data = [str(cell) for cell in row if cell is not None and str(cell).strip()]
                if row_data:
                    text += " | ".join(row_data) + "\n"
    except Exception as e:
        logger.error(f"XLSX parse error {filepath}: {e}")
    return text.strip()


def parse_document(filepath: Path) -> str:
    suffix = filepath.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(filepath)
    elif suffix in [".doc", ".docx"]:
        return parse_docx(filepath)
    elif suffix in [".xls", ".xlsx"]:
        return parse_xlsx(filepath)
    elif suffix == ".txt":
        return parse_txt(filepath)
    else:
        logger.warning(f"Unsupported file type: {filepath}")
        return ""


# ─── Ingestion Pipeline ───────────────────────────────────────────────────────

def ingest_document(filepath: Path, collection_name: str = CHROMA_COLLECTION_INTERNAL) -> dict:
    """
    Parse, chunk, embed, and store a document in ChromaDB.
    Returns a status dict.
    """
    logger.info(f"Ingesting: {filepath.name}")

    # 1. Parse
    raw_text = parse_document(filepath)
    if not raw_text:
        logger.error(f"  ✗ No text extracted from {filepath.name}")
        return {"status": "error", "file": filepath.name, "message": "Could not extract text"}

    logger.info(f"  Extracted {len(raw_text)} characters from {filepath.name}")

    # 2. Extract metadata
    metadata = extract_circular_metadata(raw_text, filepath.name)
    logger.info(f"  Metadata: Ref={metadata['ref_no']}, Date={metadata['date']}, Auth={metadata['authority']}")

    # 3. Chunk text
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ".", " ", ""],
        is_separator_regex=False
    )
    chunks = splitter.split_text(raw_text)
    logger.info(f"  Split into {len(chunks)} chunks for {filepath.name}")

    if not chunks:
        return {"status": "error", "file": filepath.name, "message": "No text chunks generated"}

    # 4. Embed
    embedder = get_embeddings()
    try:
        embeddings = embedder.embed_documents(chunks)
        logger.info(f"  Generated {len(embeddings)} embeddings (dim={len(embeddings[0]) if embeddings else '?'})")
    except Exception as e:
        logger.error(f"  ✗ Embedding failed for {filepath.name}: {e}")
        return {"status": "error", "file": filepath.name, "message": f"Embedding failed: {e}"}

    # 5. Store in ChromaDB
    collection = get_collection(collection_name)

    # Check for existing chunks from same file to avoid duplicates
    try:
        existing = collection.get(where={"filename": filepath.name})
        if existing and existing["ids"]:
            logger.info(f"  Removing {len(existing['ids'])} existing chunks for {filepath.name}")
            collection.delete(ids=existing["ids"])
    except Exception as e:
        logger.warning(f"  Could not check/remove existing chunks for {filepath.name}: {e}")
        # Continue anyway — duplicates are better than no ingestion

    ids = [str(uuid.uuid4()) for _ in chunks]
    metadatas = [{**metadata, "chunk_index": i, "total_chunks": len(chunks)} for i in range(len(chunks))]

    collection.add(
        ids=ids,
        embeddings=embeddings,
        documents=chunks,
        metadatas=metadatas
    )

    logger.info(f"  ✓ Successfully ingested {filepath.name} ({len(chunks)} chunks)")
    return {
        "status": "success",
        "file": filepath.name,
        "chunks": len(chunks),
        "ref_no": metadata["ref_no"],
        "date": metadata["date"],
        "authority": metadata["authority"],
    }


def ingest_text(text: str, metadata: dict, collection_name: str = CHROMA_COLLECTION_INTERNAL) -> dict:
    """
    Directly ingest raw text (used for web-scraped content).
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    chunks = splitter.split_text(text)
    if not chunks:
        return {"status": "error", "message": "Empty text"}

    embedder = get_embeddings()
    embeddings = embedder.embed_documents(chunks)

    from config import CHROMA_COLLECTION_EXTERNAL
    collection = get_collection(collection_name)

    ids = [str(uuid.uuid4()) for _ in chunks]
    metadatas = [{**metadata, "chunk_index": i, "total_chunks": len(chunks)} for i in range(len(chunks))]

    collection.add(
        ids=ids,
        embeddings=embeddings,
        documents=chunks,
        metadatas=metadatas
    )

    return {"status": "success", "chunks": len(chunks)}


def get_ingestion_stats() -> dict:
    """Return stats about what's stored in the vector DB."""
    try:
        client = get_chroma_client()
        stats = {}
        for name in [CHROMA_COLLECTION_INTERNAL, "regulatory_docs"]:
            try:
                col = client.get_collection(name)
                count = col.count()
                stats[name] = count
            except Exception:
                stats[name] = 0
        return stats
    except Exception as e:
        return {"error": str(e)}


def list_ingested_files(collection_name: str = CHROMA_COLLECTION_INTERNAL) -> list:
    """List all unique files that have been ingested."""
    try:
        collection = get_collection(collection_name)
        results = collection.get(include=["metadatas"])
        seen = {}
        for meta in results["metadatas"]:
            fname = meta.get("filename", "unknown")
            if fname not in seen:
                seen[fname] = {
                    "filename": fname,
                    "ref_no": meta.get("ref_no", "N/A"),
                    "date": meta.get("date", "N/A"),
                    "authority": meta.get("authority", "N/A"),
                    "ingested_at": meta.get("ingested_at", "N/A"),
                    "total_chunks": meta.get("total_chunks", 0),
                }
        return list(seen.values())
    except Exception as e:
        logger.error(f"Error listing files: {e}")
        return []


if __name__ == "__main__":
    # Quick test: ingest all files in circulars directory
    from config import CIRCULARS_DIR
    files = list(Path(CIRCULARS_DIR).glob("**/*"))
    supported = [f for f in files if f.suffix.lower() in [".pdf", ".doc", ".docx", ".xls", ".xlsx", ".txt"]]
    print(f"Found {len(supported)} documents to ingest")
    for f in supported:
        result = ingest_document(f)
        print(result)
